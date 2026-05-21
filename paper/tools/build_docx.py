#!/usr/bin/env python3
"""Build an academic Word manuscript from paper/manuscript.tex.

The LaTeX file is the canonical source. This script supports the subset used in
that manuscript: sections, subsections, abstract, itemized and enumerated lists,
basic tables, figures, display equations, and natbib-style citations.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
PAPER = ROOT / "paper"
TEX = PAPER / "manuscript.tex"
OUT = PAPER / "build" / "manuscript.docx"

CITES = {
    "argyle2023silicon": ("Argyle et al.", "2023"),
    "brand2023willingness": ("Brand et al.", "2023"),
    "gilardi2023chatgpt": ("Gilardi et al.", "2023"),
    "gema2025inverse": ("Gema et al.", "2025"),
    "hamilton2020nlp": ("Hamilton and Lahne", "2020"),
    "hamilton2025review": ("Hamilton et al.", "2026"),
    "hollmann2025tabpfn": ("Hollmann et al.", "2025"),
    "horton2023llm": ("Horton et al.", "2023"),
    "jakkaew2026coffee": ("Jakkaew et al.", "2026"),
    "kahneman2011thinking": ("Kahneman", "2011"),
    "liu2023geval": ("Liu et al.", "2023"),
    "liu2024cotperceptual": ("Liu et al.", "2024"),
    "mahieu2020fcwine": ("Mahieu et al.", "2020"),
    "mahieu2021mrca": ("Mahieu et al.", "2021"),
    "mahieu2022ifc": ("Mahieu et al.", "2022"),
    "maier2025ssr": ("Maier et al.", "2025"),
    "miller2021whisky": ("Miller et al.", "2021"),
    "motoki2025genai": ("Motoki et al.", "2025"),
    "sui2025overthinking": ("Sui et al.", "2025"),
    "symoneaux2012comments": ("Symoneaux et al.", "2012"),
    "tenkleij2003freecomment": ("ten Kleij and Musters", "2003"),
    "torrico2025chatgpt": ("Torrico", "2025"),
    "visalli2024hamdata": ("Visalli et al.", "2024a"),
    "visalli2024madeleines": ("Visalli et al.", "2025"),
    "visallimahieu2023preprocessing": ("Visalli and Mahieu", "2023"),
    "wang2025flavor": ("Wang and Pellegrino", "2025"),
    "wei2022cot": ("Wei et al.", "2022"),
    "worch2013ipm": ("Worch et al.", "2013"),
    "worchennis2013singleideal": ("Worch and Ennis", "2013"),
    "zheng2023judge": ("Zheng et al.", "2023"),
    "zhou2026reasoning": ("Zhou et al.", "2026"),
}


def citation_text(keys: str, parenthetical: bool) -> str:
    parts = []
    for key in [k.strip() for k in keys.split(",")]:
        author, year = CITES.get(key, (key, ""))
        parts.append(f"{author}, {year}" if parenthetical else f"{author} ({year})")
    sep = "; " if parenthetical else "; "
    text = sep.join(parts)
    return f"({text})" if parenthetical else text


def latex_to_text(text: str) -> str:
    text = text.replace("``", '"').replace("''", '"')
    text = re.sub(r"\\citet\{([^}]+)\}", lambda m: citation_text(m.group(1), False), text)
    text = re.sub(r"\\citep\{([^}]+)\}", lambda m: citation_text(m.group(1), True), text)
    replacements = {
        r"\Rtwo{}": "R²",
        r"\Rtwo": "R²",
        r"\mae{}": "MAE",
        r"\mae": "MAE",
        r"\fl{}": "Gemini 2.5 Flash Lite",
        r"\fl": "Gemini 2.5 Flash Lite",
        r"\gthree{}": "Gemini 3 Flash",
        r"\gthree": "Gemini 3 Flash",
        r"\%": "%",
        r"\&": "&",
        r"\_": "_",
        r"\textwidth": "",
        r"--": "-",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"\\textit\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\textbf\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\mathrm\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\texttt\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\$([^$]+)\$", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^]]+\])?(?:\{([^{}]*)\})?", lambda m: m.group(1) or "", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def equation_to_text(equation: str) -> str:
    if '"visual"' in equation and '"texture"' in equation and '"flavor"' in equation:
        return '{"visual": X, "texture": X, "flavor": X}.'
    if r"\Rtwo" in equation and r"\mae" in equation:
        return "R² = 1 - sum_i (y_i - y-hat_i)^2 / sum_i (y_i - y-bar)^2; MAE = (1/n) sum_i |y_i - y-hat_i|."
    if r"\cos" in equation and r"\mathbf" in equation:
        return "cos(a, i) = (a · i) / (||a|| ||i||)."
    return latex_to_text(equation)


def extract_braced(source: str, command: str) -> str:
    marker = f"\\{command}{{"
    start = source.find(marker)
    if start < 0:
        return ""
    i = start + len(marker)
    depth = 1
    chars = []
    while i < len(source) and depth:
        ch = source[i]
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                break
        chars.append(ch)
        i += 1
    return "".join(chars)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Title", 16), ("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    cleaned = latex_to_text(text)
    if cleaned:
        doc.add_paragraph(cleaned, style=style)


def add_title_page(doc: Document, source: str) -> None:
    title = latex_to_text(extract_braced(source, "title").replace("\\large", "").replace("\\\\", " "))
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)

    authors = [
        ("John M. Ennis", "Aigora, Richmond, Virginia, USA"),
        ("Thierry Worch", "FrieslandCampina"),
        ("Benjamin Mahieu", "Oniris Nantes"),
    ]
    for name, affiliation in authors:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(name).bold = True
        p.add_run(f"\n{affiliation}")

    date = latex_to_text(extract_braced(source, "date"))
    p = doc.add_paragraph(date)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def parse_table(doc: Document, block: str) -> None:
    caption_match = re.search(r"\\caption\{([^{}]+)\}", block, re.S)
    caption = latex_to_text(caption_match.group(1)) if caption_match else "Table"
    doc.add_paragraph(caption, style="Caption")
    begin = block.find(r"\begin{tabular}")
    end = block.find(r"\end{tabular}")
    if begin < 0 or end < 0:
        return
    body_start = block.find("\n", begin)
    if body_start < 0 or body_start >= end:
        return
    body = block[body_start:end]
    body = re.sub(r"\\(toprule|midrule|bottomrule)\s*", "", body)
    rows = []
    for raw in re.split(r"\\\\", body):
        raw = raw.strip()
        if not raw:
            continue
        rows.append([latex_to_text(cell.strip()) for cell in re.split(r"(?<!\\)&", raw)])
    if not rows:
        return
    n_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c in range(n_cols):
            table.cell(r, c).text = row[c] if c < len(row) else ""
            for paragraph in table.cell(r, c).paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def parse_figure(doc: Document, block: str) -> None:
    include_match = re.search(r"\\includegraphics(?:\[[^]]+\])?\{([^}]+)\}", block)
    caption_match = re.search(r"\\caption\{([^{}]+)\}", block, re.S)
    caption = latex_to_text(caption_match.group(1)) if caption_match else "Figure"
    if include_match:
        image_name = include_match.group(1)
        image_path = PAPER / "figures" / image_name
        if not image_path.exists() and image_path.suffix.lower() == ".pdf":
            image_path = image_path.with_suffix(".png")
        if image_path.exists():
            doc.add_picture(str(image_path), width=Inches(5.9))
    doc.add_paragraph(caption, style="Caption")


def parse_bibliography(doc: Document, block: str) -> None:
    doc.add_heading("References", level=1)
    entries = re.split(r"\\bibitem(?:\[[^]]+\])?\{[^}]+\}", block)
    for entry in entries[1:]:
        entry = re.sub(r"\\end\{thebibliography\}", "", entry)
        text = latex_to_text(entry)
        if text:
            p = doc.add_paragraph(text)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.left_indent = Inches(0.25)


def build_docx() -> None:
    source = TEX.read_text(encoding="utf-8")
    body_match = re.search(r"\\begin\{document\}(.*)\\end\{document\}", source, re.S)
    if not body_match:
        raise RuntimeError("Could not find document body")
    body = body_match.group(1)

    doc = Document()
    configure_document(doc)
    add_title_page(doc, source)

    pos = 0
    paragraph_lines: list[str] = []
    list_kind: str | None = None

    token_re = re.compile(
        r"\\begin\{abstract\}.*?\\end\{abstract\}|"
        r"\\begin\{table\}.*?\\end\{table\}|"
        r"\\begin\{figure\}.*?\\end\{figure\}|"
        r"\\begin\{thebibliography\}.*?\\end\{thebibliography\}|"
        r"\\\[.*?\\\]",
        re.S,
    )

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            add_paragraph(doc, " ".join(paragraph_lines))
            paragraph_lines = []

    def handle_line(line: str) -> None:
        nonlocal list_kind
        stripped = line.strip()
        if not stripped or stripped in {r"\maketitle", r"\centering"}:
            flush_paragraph()
            return
        section = re.match(r"\\section\*?\{([^}]+)\}", stripped)
        if section:
            flush_paragraph()
            doc.add_heading(latex_to_text(section.group(1)), level=1)
            return
        subsection = re.match(r"\\subsection\{([^}]+)\}", stripped)
        if subsection:
            flush_paragraph()
            doc.add_heading(latex_to_text(subsection.group(1)), level=2)
            return
        subsubsection = re.match(r"\\subsubsection\{([^}]+)\}", stripped)
        if subsubsection:
            flush_paragraph()
            doc.add_heading(latex_to_text(subsubsection.group(1)), level=3)
            return
        if stripped == r"\begin{itemize}":
            flush_paragraph()
            list_kind = "List Bullet"
            return
        if stripped == r"\begin{enumerate}":
            flush_paragraph()
            list_kind = "List Number"
            return
        if stripped in {r"\end{itemize}", r"\end{enumerate}"}:
            flush_paragraph()
            list_kind = None
            return
        item = re.match(r"\\item\s+(.+)", stripped)
        if item and list_kind:
            flush_paragraph()
            add_paragraph(doc, item.group(1), style=list_kind)
            return
        if stripped.startswith("\\label"):
            return
        paragraph_lines.append(stripped)

    for match in token_re.finditer(body):
        for line in body[pos : match.start()].splitlines():
            handle_line(line)
        flush_paragraph()
        block = match.group(0)
        if block.startswith(r"\begin{abstract}"):
            doc.add_heading("Abstract", level=1)
            text = re.sub(r"\\begin\{abstract\}|\\end\{abstract\}", "", block)
            add_paragraph(doc, text)
        elif block.startswith(r"\begin{table}"):
            parse_table(doc, block)
        elif block.startswith(r"\begin{figure}"):
            parse_figure(doc, block)
        elif block.startswith(r"\begin{thebibliography}"):
            parse_bibliography(doc, block)
        elif block.startswith(r"\["):
            equation = block.replace(r"\[", "").replace(r"\]", "").strip()
            p = doc.add_paragraph(equation_to_text(equation))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pos = match.end()

    for line in body[pos:].splitlines():
        handle_line(line)
    flush_paragraph()

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build_docx()
